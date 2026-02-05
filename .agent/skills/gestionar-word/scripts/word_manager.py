import argparse
import sys
import json
import os

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Error: 'python-docx' library is not installed. Please install it using 'pip install python-docx'.", file=sys.stderr)
    sys.exit(1)

def read_document(file_path):
    """Reads the full text of the document."""
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading document: {e}"

def inspect_document(file_path):
    """Inspects the document structure, including bold text and comments (if accessible via XML)."""
    try:
        doc = Document(file_path)
        inspection_data = {
            "paragraphs": [],
            "comments": [] # Comments are complex in python-docx, often requiring XML parsing.
        }

        for i, para in enumerate(doc.paragraphs):
            para_data = {
                "index": i,
                "text": para.text,
                "bold_runs": []
            }
            for run in para.runs:
                if run.bold:
                    para_data["bold_runs"].append(run.text)
            
            if para_data["bold_runs"]:
                 inspection_data["paragraphs"].append(para_data)

        # Basic XML parsing for comments (w:comments) if reliable way exists, 
        # but for now we focus on content structure as python-docx minimal support for comments is tricky.
        # We will iterate to see if we can extract more if requested.
        
        return json.dumps(inspection_data, indent=2, ensure_ascii=False)
    except Exception as e:
        return f"Error inspecting document: {e}"

def write_document(file_path, content, mode='a'):
    """Writes content to a document. Mode 'a' for append, 'w' for overwrite (create new)."""
    try:
        if mode == 'w' or not os.path.exists(file_path):
            doc = Document()
        else:
            doc = Document(file_path)
        
        doc.add_paragraph(content)
        doc.save(file_path)
        return f"Successfully wrote to {file_path}"
    except Exception as e:
        return f"Error writing to document: {e}"

def create_document(file_path, content=""):
    """Creates a new document with optional initial content."""
    return write_document(file_path, content, mode='w')

def main():
    parser = argparse.ArgumentParser(description="Word Document Manager for Antigravity Agent")
    subparsers = parser.add_subparsers(dest="command", help="Command to execute")

    # Read command
    read_parser = subparsers.add_parser("read", help="Read text from a document")
    read_parser.add_argument("file_path", help="Path to the .docx file")

    # Inspect command
    inspect_parser = subparsers.add_parser("inspect", help="Inspect document structure (bold, etc.)")
    inspect_parser.add_argument("file_path", help="Path to the .docx file")

    # Write command
    write_parser = subparsers.add_parser("write", help="Append text to a document")
    write_parser.add_argument("file_path", help="Path to the .docx file")
    write_parser.add_argument("content", help="Content to write")

    # Create command
    create_parser = subparsers.add_parser("create", help="Create a new document")
    create_parser.add_argument("file_path", help="Path to the new .docx file")
    create_parser.add_argument("content", nargs='?', default="", help="Initial content (optional)")

    args = parser.parse_args()

    if args.command == "read":
        print(read_document(args.file_path))
    elif args.command == "inspect":
        print(inspect_document(args.file_path))
    elif args.command == "write":
        print(write_document(args.file_path, args.content))
    elif args.command == "create":
        print(create_document(args.file_path, args.content))
    else:
        parser.print_help()

if __name__ == "__main__":
    main()
